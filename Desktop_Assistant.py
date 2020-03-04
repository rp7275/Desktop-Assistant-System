import pyttsx3
import speech_recognition as sr
import datetime
import wikipedia
import webbrowser
import os
import smtplib
import re
import sys

from docx import Document
from docx.shared import Inches





engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
# print(voices[1].id)
engine.setProperty('voice', voices[0].id)


def speak(audio):
    print(audio)
    engine.say(audio)
    engine.runAndWait()


def wishMe():
    hour = int(datetime.datetime.now().hour)
    if hour>=0 and hour<12:
        speak("Good Morning!")

    elif hour>=12 and hour<18:
        speak("Good Afternoon!")

    else:
        speak("Good Evening!")

    speak("I am Jarvis Sir. Please tell me how may I help you")

def takeCommand():
    #It takes microphone input from the user and returns string output
    
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        r.pause_threshold = 1
        r.adjust_for_ambient_noise(source, duration=1)
        audio = r.listen(source)

    try:
        print("Recognizing...")
        query = r.recognize_google(audio, language='en-in')
        print(f"User said: {query}\n")

    except Exception as e:
        # print(e)
        print("Say ehat again please...")
        return "None"
    return query

def sendEmail(to, content):
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.ehlo()
    server.starttls()
    server.login('youremail@gmail.com', 'your-password')
    server.sendmail('youremail@gmail.com', to, content)
    server.close()
    

def create_word():
    speak("Sir Give Document Title")
    query = takeCommand().lower()
    
    document = Document()
    document.add_heading(query, 0)
    speak("Sir please Give the heading ")
    query = takeCommand().lower()
    
    document.add_heading(query, level=1)
    speak("Sir please speak the contents of the document")
    query = takeCommand().lower()
    
    document.add_paragraph(query, style='Intense Quote')    
    document.add_page_break()
    speak("Give the file name with which you would like to save the document")
    query = takeCommand().lower()
    
    document.save(query +'.docx')
        
        
        

def Run_jarvis():
   
   if __name__ == "__main__":
    wishMe()
    while True:
    # if 1:
        query = takeCommand().lower()

        # Logic for executing tasks based on query
        if 'go to' in query:
            reg_ex = re.search('go to (.+)', query)
            if reg_ex:
                domain = reg_ex.group(1)
                print(domain)
                url = 'https://www.' + domain
                webbrowser.open(url)
                speak('The website you have requested has been opened for you Sir.')
            else:
                    pass


        elif 'open control panel' in query:
             os.system("cmd /c control")

        elif 'open control access' in query:
              os.system("cmd /c control access.cpl")

        elif'close' in query:
         reg_ex = re.search('close (.+)', query)
         if reg_ex:
                domain = reg_ex.group(1)
                print(domain)
                url = 'cmd /c TASKKILL /IM ' + domain+'.exe'
                print(url)
                os.system(url)
                #speak('The website you have requested has been opened for you Sir.')
         else:
                    pass


        elif 'open add or remove programs' in query:
              os.system("cmd /c control appwiz.cpl")

        elif 'open backup and restore Center' in query:
              os.system("cmd /c control /name Microsoft.BackupAndRestoreCenter")

        elif 'change date and time' in query:
              os.system("cmd /c control /name Microsoft.DateAndTime")

        elif 'change default programs' in query:
              os.system("cmd /c control /name Microsoft.DefaultPrograms")

        elif 'open device manager' in query:
              os.system("cmd /c control /name Microsoft.DeviceManager")

        elif 'show devices and printers' in query:
              os.system("cmd /c control /name Microsoft.DevicesAndPrinters")

        elif 'open display settings' in query:
              os.system("cmd /c control /name Microsoft.Display")

        elif 'show file history' in query:
            os.system("cmd /c control /name Microsoft.FileHistory")

        elif 'open folder options' in query:
            os.system("cmd /c control /name Microsoft.FolderOptions")

        elif 'change fonts' in query:
            os.system("cmd /c control /name Microsoft.Fonts")

        elif 'open game controllers' in query:
            os.system("cmd /c control /name Microsoft.GameControllers")

        elif 'open keyboard properties' in query:
            os.system("cmd /c control /name Microsoft.Keyboard")

        elif 'open language settings' in query:
            os.system("cmd /c control /name Microsoft.Language")

        elif 'open location settings' in query:
            os.system("cmd /c control /name Microsoft.LocationAndOtherSensors")

        elif 'open mouse settings' in query:
            os.system("cmd /c control /name Microsoft.Mouse")

        elif 'open network and sharing ' in query:
            os.system("cmd /c control /name Microsoft.NetworkAndSharingCenter")

        elif 'open network connections' in query:
            os.system("cmd /c control ncpa.cpl")

        elif 'open personalization' in query:
            os.system("cmd /c control /name Microsoft.Personalization")

        elif 'open power settings' in query:
            os.system("cmd /c control /name Microsoft.PowerOptions")
        elif 'create document' in query:
            
            speak("Sir Give Document Title")
            query = takeCommand().lower()
            document = Document()
            document.add_heading(query, 0)
            speak("Sir please Give the heading ")
            query = takeCommand().lower()
    
            document.add_heading(query, level=1)
            speak("Sir please speak the contents of the document")
            query = takeCommand().lower()
    
            document.add_paragraph(query, style='Intense Quote')    
            document.add_page_break()
            speak("Give the file name with which you would like to save the document")
            query = takeCommand().lower()
    
            document.save(query +'.docx')

        elif 'open security and maintenance' in query:
            os.system("cmd /c control /name Microsoft.ProblemReportsAndSolutions")

        elif 'open recovery settings' in query:
            os.system("cmd /c control /name Microsoft.Recovery")

        elif 'open remote app and desktop ' in query:
            os.system("cmd /c control /name Microsoft.RemoteAppAndDesktopConnections")

        elif 'open user accounts' in query:
            os.system("cmd /c control /name Microsoft.UserAccounts")

        elif 'open troubleshooting' in query:
            os.system("cmd /c control /name Microsoft.Troubleshooting")

        elif 'open taskbar settings' in query:
            os.system("cmd /c control /name Microsoft.TaskbarAndStartMenu")

        elif 'open system properties' in query:
            os.system("cmd /c control /name Microsoft.System")

        elif 'open system properties' in query:
            os.system("cmd /c control /name Microsoft.System")

        elif 'play music' in query:
            music_dir = 'E:\\Songs\\Billboard Hot 100 Singles Chart (7th Oct 2017)'
            songs = os.listdir(music_dir)
            print(songs)
            os.startfile(os.path.join(music_dir, songs[0]))

        elif 'the time' in query:
            strTime = datetime.datetime.now().strftime("%H hours %M minutes and %S seconds")
            speak("Sir, the time is" + strTime)

        elif 'open code' in query:
            codePath = "C:\\Users\\Dhruv\\AppData\\Local\\Programs\\Microsoft VS Code\\Code.exe"
            os.startfile(codePath)

        elif 'email to rp' in query:
            try:
                speak("What should I say?")
                content = takeCommand()
                to = "rp@gmail.com"
                sendEmail(to, content)
                speak("Email has been sent!")
            except Exception as e:
                print(e)
                speak("Sorry my friend RP. I am not able to send this email")
        
        elif 'shutdown' in query:
            sys.exit()
Run_jarvis()